#!/usr/bin/env python3
"""Extract text from Word docs and compare with markdown sources.

Usage:
    python docx_reconcile.py <docx_file> <md_file>
    python docx_reconcile.py --batch <docx_dir>

This tool helps reconcile changes made in published Word documents
back to the source markdown files.
"""

from __future__ import annotations

import argparse
import difflib
import re
import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Known mappings from Word doc names to .md sources
DOC_MAPPINGS = {
    "reactor-ops-log": "docs/requirements/prd_reactor-ops-log.md",
    "ops-log": "docs/requirements/prd_reactor-ops-log.md",
    "experiment-manager": "docs/requirements/prd_experiment-manager.md",
    "master-prd": "docs/requirements/prd_neutron-os-executive.md",
    "executive-prd": "docs/requirements/prd_neutron-os-executive.md",
    "neutron-os-executive": "docs/requirements/prd_neutron-os-executive.md",
    "data-architecture": "docs/specs/data-architecture-spec.md",
}


def extract_docx_text(docx_path: Path) -> str:
    """Extract plain text from a Word document."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(docx_path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Detect headings by style
            style_name = (para.style.name or "") if para.style else ""
            if "Heading" in style_name:
                # Convert to markdown heading
                level = 1
                if "1" in style_name:
                    level = 1
                elif "2" in style_name:
                    level = 2
                elif "3" in style_name:
                    level = 3
                elif "4" in style_name:
                    level = 4
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def normalize_text(text: str) -> str:
    """Normalize text for comparison (remove formatting noise)."""
    # Remove multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    # Remove trailing whitespace
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    return text.strip()


def extract_md_text(md_path: Path) -> str:
    """Extract text from markdown, stripping formatting for comparison."""
    content = md_path.read_text(encoding="utf-8")

    # Remove code blocks (preserve content structure but simplify)
    content = re.sub(r'```[\s\S]*?```', '[CODE BLOCK]', content)

    # Remove inline code
    content = re.sub(r'`[^`]+`', lambda m: m.group(0).strip('`'), content)

    # Remove link syntax but keep text
    content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)

    # Remove bold/italic markers
    content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
    content = re.sub(r'\*([^*]+)\*', r'\1', content)

    return normalize_text(content)


def show_diff(md_text: str, docx_text: str, context_lines: int = 3) -> list[str]:
    """Generate a unified diff between markdown and docx content."""
    md_lines = md_text.split('\n')
    docx_lines = docx_text.split('\n')

    diff = difflib.unified_diff(
        md_lines,
        docx_lines,
        fromfile='source.md',
        tofile='published.docx',
        lineterm='',
        n=context_lines,
    )

    return list(diff)


def find_md_source(docx_name: str, root: Path) -> Optional[Path]:
    """Find the corresponding .md source for a docx file."""
    docx_stem = Path(docx_name).stem.lower()

    # Try exact mappings
    for key, md_path in DOC_MAPPINGS.items():
        if key in docx_stem:
            full_path = root / md_path
            if full_path.exists():
                return full_path

    # Try fuzzy match on docs/requirements/ and docs/specs/
    for search_dir in [root / "docs/requirements", root / "docs/specs"]:
        if search_dir.exists():
            for md_file in search_dir.glob("*.md"):
                if docx_stem in md_file.stem.lower() or md_file.stem.lower() in docx_stem:
                    return md_file

    return None


def reconcile_file(docx_path: Path, md_path: Optional[Path], root: Path) -> dict:
    """Reconcile a single docx file with its markdown source."""
    result = {
        "docx": str(docx_path),
        "md": None,
        "status": "unknown",
        "diff_lines": 0,
        "additions": 0,
        "deletions": 0,
        "diff": [],
    }

    # Find md source if not provided
    if md_path is None:
        md_path = find_md_source(docx_path.name, root)

    if md_path is None or not md_path.exists():
        result["status"] = "no_source"
        return result

    result["md"] = str(md_path)

    # Extract text
    try:
        docx_text = normalize_text(extract_docx_text(docx_path))
        md_text = extract_md_text(md_path)
    except Exception as e:
        result["status"] = "error"
        result["error"] = str(e)
        return result

    # Generate diff
    diff = show_diff(md_text, docx_text)
    result["diff"] = diff
    result["diff_lines"] = len(diff)

    # Count additions/deletions
    for line in diff:
        if line.startswith('+') and not line.startswith('+++'):
            result["additions"] += 1
        elif line.startswith('-') and not line.startswith('---'):
            result["deletions"] += 1

    if result["additions"] == 0 and result["deletions"] == 0:
        result["status"] = "in_sync"
    else:
        result["status"] = "diverged"

    return result


def main():
    parser = argparse.ArgumentParser(description="Reconcile Word docs with markdown sources")
    parser.add_argument("docx_path", help="Path to .docx file or directory")
    parser.add_argument("md_path", nargs="?", help="Path to corresponding .md file (auto-detected if omitted)")
    parser.add_argument("--root", default=".", help="Repository root (default: current directory)")
    parser.add_argument("--batch", action="store_true", help="Process all .docx files in directory")
    parser.add_argument("--output", "-o", help="Output diff to file")
    parser.add_argument("--summary", action="store_true", help="Show summary only, not full diff")
    args = parser.parse_args()

    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    root = Path(args.root).resolve()
    docx_path = Path(args.docx_path)

    if args.batch or docx_path.is_dir():
        # Batch mode: process all .docx in directory
        docx_dir = docx_path if docx_path.is_dir() else docx_path.parent
        docx_files = list(docx_dir.glob("*.docx"))

        if not docx_files:
            print(f"No .docx files found in {docx_dir}")
            sys.exit(1)

        print(f"Processing {len(docx_files)} Word documents...\n")

        for docx_file in sorted(docx_files):
            result = reconcile_file(docx_file, None, root)
            status_icon = {
                "in_sync": "✓",
                "diverged": "⚠",
                "no_source": "?",
                "error": "✗",
            }.get(result["status"], "?")

            print(f"{status_icon} {docx_file.name}")
            if result["md"]:
                print(f"  → {result['md']}")
            if result["status"] == "diverged":
                print(f"  Changes: +{result['additions']} -{result['deletions']} lines")
            elif result["status"] == "no_source":
                print("  No matching .md source found")
            elif result["status"] == "error":
                print(f"  Error: {result.get('error', 'unknown')}")
            print()
    else:
        # Single file mode
        md_path = Path(args.md_path) if args.md_path else None
        result = reconcile_file(docx_path, md_path, root)

        print(f"Document: {result['docx']}")
        print(f"Source:   {result['md'] or 'Not found'}")
        print(f"Status:   {result['status']}")
        print()

        if result["status"] == "diverged":
            print(f"Changes: +{result['additions']} additions, -{result['deletions']} deletions\n")

            if not args.summary:
                print("=" * 60)
                print("DIFF (- = in .md only, + = in .docx only)")
                print("=" * 60)
                for line in result["diff"]:
                    print(line)

        if args.output and result["diff"]:
            Path(args.output).write_text('\n'.join(result["diff"]))
            print(f"\nDiff written to: {args.output}")


if __name__ == "__main__":
    main()
