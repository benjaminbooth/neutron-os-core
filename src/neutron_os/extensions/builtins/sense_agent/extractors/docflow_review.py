"""DocFlow Review Extractor — imports external document changes as signals.

Handles bidirectional sync between git .md files and external documents
(MS 365 Word, Google Docs, etc.) where external docs are authoritative
after publication.

Key insight: External doc changes are just another signal source.
Reviewer comments, tracked changes, and content drift all become
signals in the unified suggest queue.

Reconciliation strategy:
1. DETECT divergence between .md and external doc
2. EXTRACT changes/comments from external doc as signals
3. If both diverged: LLM-assisted semantic merge (external wins by default)
4. Generate PRD suggestions with high confidence (95%) for review changes

Supported formats:
- MS 365 Word (.docx) via python-docx + Graph API
- Google Docs via Google Docs API
- Local .docx files (offline mode)
"""

from __future__ import annotations

import hashlib
import json
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Optional

from ..models import Signal


class DocFormat(Enum):
    """Supported external document formats."""
    MS_WORD = "ms_word"
    GOOGLE_DOCS = "google_docs"
    LOCAL_DOCX = "local_docx"


class ChangeType(Enum):
    """Types of changes detected in external documents."""
    COMMENT = "comment"           # Reviewer comment
    TRACKED_CHANGE = "tracked"    # Track changes addition/deletion
    CONTENT_DRIFT = "drift"       # Content differs from .md (no tracking)
    STRUCTURAL = "structural"     # Section added/removed/reordered


@dataclass
class ExternalChange:
    """A change detected in an external document."""
    change_type: ChangeType
    section: str                  # Which PRD section (e.g., "Requirements")
    author: str                   # Who made the change
    timestamp: str                # When (if available)
    original_text: str            # Text in .md
    new_text: str                 # Text in external doc
    comment_text: str = ""        # If COMMENT type, the actual comment
    context: str = ""             # Surrounding text for LLM reasoning
    confidence: float = 0.95      # External doc changes = high confidence


@dataclass
class DivergenceReport:
    """Report of divergence between .md and external document."""
    prd_id: str
    md_path: Path
    external_uri: str             # SharePoint URL, Google Doc ID, etc.
    external_format: DocFormat

    # Content hashes
    md_hash: str
    external_hash: str
    last_sync_hash: str = ""      # Hash at last known sync point

    # Divergence analysis
    md_changed_since_sync: bool = False
    external_changed_since_sync: bool = False
    requires_merge: bool = False  # Both diverged = need LLM merge

    # Detected changes
    changes: list[ExternalChange] = field(default_factory=list)

    # Merge result (if merge was needed)
    merge_strategy: str = ""      # "external_wins", "llm_merge", "conflict"
    merged_content: str = ""


class DocProvider(ABC):
    """Abstract base for document providers (MS 365, Google, local)."""

    @abstractmethod
    def fetch_content(self, uri: str) -> tuple[str, list[ExternalChange]]:
        """Fetch document content and extract changes.

        Returns:
            Tuple of (plain text content, list of changes/comments)
        """
        pass

    @abstractmethod
    def get_revision_history(self, uri: str) -> list[dict]:
        """Get document revision history for merge analysis."""
        pass


class LocalDocxProvider(DocProvider):
    """Provider for local .docx files (offline development)."""

    def __init__(self):
        try:
            import docx
            self._docx = docx
        except ImportError:
            self._docx = None

    def fetch_content(self, uri: str) -> tuple[str, list[ExternalChange]]:
        """Extract text and comments from local .docx file."""
        if not self._docx:
            raise ImportError("python-docx required: pip install python-docx")

        path = Path(uri)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {uri}")

        doc = self._docx.Document(str(path))
        changes: list[ExternalChange] = []

        # Extract plain text
        paragraphs = [p.text for p in doc.paragraphs]
        content = "\n".join(paragraphs)

        # Extract comments (if any)
        # Note: python-docx has limited comment support;
        # full support requires parsing the XML directly
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            comments_part = doc.part.related_parts.get(RT.COMMENTS)
            if comments_part:
                # Parse comments XML
                # This is simplified; real impl would extract full comment data
                pass
        except Exception:
            pass  # Comments extraction failed, continue without

        return content, changes

    def get_revision_history(self, uri: str) -> list[dict]:
        """Local files don't have revision history."""
        return []


class MSGraphProvider(DocProvider):
    """Provider for MS 365 documents via Graph API."""

    def __init__(self, access_token: Optional[str] = None):
        self._token = access_token
        self._base_url = "https://graph.microsoft.com/v1.0"

    def fetch_content(self, uri: str) -> tuple[str, list[ExternalChange]]:
        """Fetch Word doc content and comments via Graph API.

        Args:
            uri: SharePoint/OneDrive document URL or drive item ID
        """
        if not self._token:
            raise ValueError("MS Graph access token required")

        import requests

        headers = {"Authorization": f"Bearer {self._token}"}
        changes: list[ExternalChange] = []

        # Convert SharePoint URL to drive item if needed
        item_id = self._resolve_item_id(uri)

        # Get document content (as plain text)
        # Note: For .docx, we'd need to download and parse
        # Graph API can convert to PDF/HTML but not plain text directly

        # Get comments
        comments_url = f"{self._base_url}/me/drive/items/{item_id}/comments"
        resp = requests.get(comments_url, headers=headers)
        if resp.ok:
            for comment in resp.json().get("value", []):
                changes.append(ExternalChange(
                    change_type=ChangeType.COMMENT,
                    section="",  # Would need to correlate with position
                    author=comment.get("author", {}).get("displayName", "Unknown"),
                    timestamp=comment.get("createdDateTime", ""),
                    original_text="",
                    new_text="",
                    comment_text=comment.get("content", ""),
                    confidence=0.95,
                ))

        # Download and parse .docx for full content
        content = self._download_and_parse(item_id, headers)

        return content, changes

    def _resolve_item_id(self, uri: str) -> str:
        """Convert SharePoint URL to drive item ID."""
        # Implementation would parse SharePoint URLs
        # For now, assume uri is already an item ID
        return uri

    def _download_and_parse(self, item_id: str, headers: dict) -> str:
        """Download .docx and extract plain text."""
        import io
        import requests

        try:
            import docx
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        url = f"{self._base_url}/me/drive/items/{item_id}/content"
        resp = requests.get(url, headers=headers)
        resp.raise_for_status()

        doc = docx.Document(io.BytesIO(resp.content))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)

    def get_revision_history(self, uri: str) -> list[dict]:
        """Get document version history from OneDrive/SharePoint."""
        if not self._token:
            return []

        import requests

        headers = {"Authorization": f"Bearer {self._token}"}
        item_id = self._resolve_item_id(uri)

        url = f"{self._base_url}/me/drive/items/{item_id}/versions"
        resp = requests.get(url, headers=headers)

        if resp.ok:
            return resp.json().get("value", [])
        return []


class DocFlowReviewExtractor:
    """Extracts signals from external document changes.

    Workflow:
    1. Load PRD registry (maps prd_id → .md path + external doc URI)
    2. For each registered PRD with external doc:
       a. Fetch external content
       b. Compare with local .md
       c. Detect divergence
       d. Extract changes as signals
       e. If both diverged: trigger LLM merge
    3. Output signals to unified suggest queue
    """

    def __init__(
        self,
        registry_path: Optional[Path] = None,
        llm_gateway = None,
    ):
        from neutron_os import REPO_ROOT as _REPO_ROOT
        self._runtime_dir = _REPO_ROOT / "runtime"
        self._registry_path = registry_path or (
            self._runtime_dir / "inbox" / "state" / "docflow_registry.json"
        )
        self._sync_state_path = self._runtime_dir / "inbox" / "state" / "docflow_sync.json"
        self._gateway = llm_gateway

        # Initialize legacy providers (for backward compatibility)
        self._providers: dict[DocFormat, DocProvider] = {
            DocFormat.LOCAL_DOCX: LocalDocxProvider(),
        }

        # Try to initialize MS Graph provider
        ms_token = self._get_ms_token()
        if ms_token:
            self._providers[DocFormat.MS_WORD] = MSGraphProvider(ms_token)

        # Import new provider registry (if available)
        try:
            from .docflow_providers import (
                ProviderRegistry,
                sync_all_folders,
            )
            self._provider_registry = ProviderRegistry
            self._sync_all_folders = sync_all_folders
        except ImportError:
            self._provider_registry = None
            self._sync_all_folders = None

    def _get_ms_token(self) -> Optional[str]:
        """Get MS Graph access token from environment or token cache."""
        import os
        return os.environ.get("MS_GRAPH_TOKEN")

    def _load_registry(self) -> dict:
        """Load PRD ↔ external doc registry."""
        if not self._registry_path.exists():
            return {"prds": {}}
        return json.loads(self._registry_path.read_text())

    def _load_sync_state(self) -> dict:
        """Load last-known sync state (hashes at sync point)."""
        if not self._sync_state_path.exists():
            return {}
        return json.loads(self._sync_state_path.read_text())

    def _save_sync_state(self, state: dict) -> None:
        """Save sync state after successful sync."""
        self._sync_state_path.parent.mkdir(parents=True, exist_ok=True)
        self._sync_state_path.write_text(json.dumps(state, indent=2))

    def _hash_content(self, content: str) -> str:
        """Compute content hash for change detection."""
        return hashlib.sha256(content.encode()).hexdigest()[:16]

    def _read_md(self, path: Path) -> str:
        """Read .md file content."""
        if not path.exists():
            return ""
        return path.read_text(encoding="utf-8")

    def analyze_divergence(
        self,
        prd_id: str,
        md_path: Path,
        external_uri: str,
        doc_format: DocFormat,
    ) -> DivergenceReport:
        """Analyze divergence between .md and external document.

        Returns detailed report of what changed and whether merge is needed.
        """
        provider = self._providers.get(doc_format)
        if not provider:
            raise ValueError(f"No provider for format: {doc_format}")

        # Read local .md
        md_content = self._read_md(md_path)
        md_hash = self._hash_content(md_content)

        # Fetch external content
        external_content, external_changes = provider.fetch_content(external_uri)
        external_hash = self._hash_content(external_content)

        # Load sync state
        sync_state = self._load_sync_state()
        last_sync = sync_state.get(prd_id, {})
        last_sync_hash = last_sync.get("hash", "")

        # Determine what changed
        md_changed = md_hash != last_sync_hash if last_sync_hash else False
        external_changed = external_hash != last_sync_hash if last_sync_hash else True
        requires_merge = md_changed and external_changed

        # Detect content drift (when no tracked changes)
        if external_hash != md_hash and not external_changes:
            drift_changes = self._detect_content_drift(md_content, external_content)
            external_changes.extend(drift_changes)

        return DivergenceReport(
            prd_id=prd_id,
            md_path=md_path,
            external_uri=external_uri,
            external_format=doc_format,
            md_hash=md_hash,
            external_hash=external_hash,
            last_sync_hash=last_sync_hash,
            md_changed_since_sync=md_changed,
            external_changed_since_sync=external_changed,
            requires_merge=requires_merge,
            changes=external_changes,
        )

    def _detect_content_drift(
        self,
        md_content: str,
        external_content: str,
    ) -> list[ExternalChange]:
        """Detect content differences when there's no tracked changes.

        Uses difflib for structural comparison, then LLM for semantic analysis.
        """
        import difflib

        changes: list[ExternalChange] = []

        # Simple line-by-line diff
        md_lines = md_content.splitlines()
        ext_lines = external_content.splitlines()

        differ = difflib.unified_diff(md_lines, ext_lines, lineterm="")
        diff_text = "\n".join(differ)

        if diff_text:
            # For now, create a single CONTENT_DRIFT change
            # LLM will break this down into semantic sections
            changes.append(ExternalChange(
                change_type=ChangeType.CONTENT_DRIFT,
                section="[multiple]",
                author="external_doc",
                timestamp=datetime.now(timezone.utc).isoformat(),
                original_text=md_content[:500] + "..." if len(md_content) > 500 else md_content,
                new_text=external_content[:500] + "..." if len(external_content) > 500 else external_content,
                context=diff_text[:1000],
                confidence=0.90,  # Drift detection = slightly lower confidence
            ))

        return changes

    def reconcile_with_llm(
        self,
        report: DivergenceReport,
    ) -> str:
        """Use LLM to perform semantic merge when both versions diverged.

        Strategy:
        1. External doc is authoritative (it has human review feedback)
        2. But .md may have valid additions (new technical details)
        3. LLM identifies what to keep from each, avoiding duplication
        4. Result is conceptually coherent merged content
        """
        if not self._gateway:
            raise ValueError("LLM gateway required for merge reconciliation")

        md_content = self._read_md(report.md_path)

        # For external content, we need to re-fetch
        provider = self._providers.get(report.external_format)
        if not provider:
            raise ValueError(f"No provider for {report.external_format}")

        external_content, _ = provider.fetch_content(report.external_uri)

        prompt = f"""You are reconciling two versions of a Product Requirements Document (PRD).

## Context
- The EXTERNAL version (.docx in MS 365) has been reviewed and edited by stakeholders
- The LOCAL version (.md in git) may have technical additions from developers
- EXTERNAL is authoritative for requirements, scope, and stakeholder decisions
- LOCAL may have valid technical details that should be preserved

## External Version (AUTHORITATIVE)
```markdown
{external_content}
```

## Local Version (.md)
```markdown
{md_content}
```

## Instructions
1. Start with the EXTERNAL version as the base
2. Identify any technical additions in LOCAL that don't conflict with EXTERNAL
3. Merge those additions into appropriate sections
4. If there's a direct conflict (same topic, different content), EXTERNAL wins
5. Ensure the result is coherent and doesn't have duplicate content

## Output
Return ONLY the merged markdown content, no explanations.
"""

        response = self._gateway.complete(prompt, max_tokens=4000)
        report.merge_strategy = "llm_merge"
        report.merged_content = response.content

        return response.content

    def extract(
        self,
        gateway = None,
        correlator = None,
    ) -> list[Signal]:
        """Extract signals from all registered external documents.

        Returns:
            List of Signal objects for the suggest queue
        """
        if gateway:
            self._gateway = gateway

        signals: list[Signal] = []
        registry = self._load_registry()

        for prd_id, prd_config in registry.get("prds", {}).items():
            md_path = Path(prd_config.get("md_path", ""))
            external_uri = prd_config.get("external_uri", "")
            format_str = prd_config.get("format", "local_docx")

            try:
                doc_format = DocFormat(format_str)
            except ValueError:
                continue

            if not md_path or not external_uri:
                continue

            try:
                report = self.analyze_divergence(
                    prd_id=prd_id,
                    md_path=md_path,
                    external_uri=external_uri,
                    doc_format=doc_format,
                )

                # Convert changes to signals
                for change in report.changes:
                    signal = Signal(
                        source="docflow_review",
                        timestamp=change.timestamp or datetime.now(timezone.utc).isoformat(),
                        raw_text=change.comment_text or change.new_text,
                        people=[change.author] if change.author else [],
                        initiatives=[prd_id],
                        signal_type=self._change_type_to_signal_type(change.change_type),
                        detail=self._format_change_detail(change),
                        confidence=change.confidence,
                        metadata={
                            "change_type": change.change_type.value,
                            "section": change.section,
                            "external_uri": external_uri,
                            "requires_merge": report.requires_merge,
                        },
                    )
                    signals.append(signal)

                # If merge needed, add a special merge signal
                if report.requires_merge:
                    signals.append(Signal(
                        source="docflow_review",
                        timestamp=datetime.now(timezone.utc).isoformat(),
                        raw_text="Divergence detected: both .md and external doc changed since last sync",
                        initiatives=[prd_id],
                        signal_type="action_item",
                        detail=f"PRD {prd_id} requires merge reconciliation. .md and external doc have both been modified.",
                        confidence=1.0,
                        metadata={
                            "requires_merge": True,
                            "md_hash": report.md_hash,
                            "external_hash": report.external_hash,
                            "external_uri": external_uri,
                        },
                    ))

            except Exception as e:
                # Log error but continue with other PRDs
                signals.append(Signal(
                    source="docflow_review",
                    timestamp=datetime.now(timezone.utc).isoformat(),
                    raw_text=f"Failed to analyze {prd_id}: {str(e)}",
                    initiatives=[prd_id],
                    signal_type="blocker",
                    detail=f"DocFlow extraction failed for {prd_id}",
                    confidence=0.5,
                    metadata={"error": str(e)},
                ))

        return signals

    def _change_type_to_signal_type(self, change_type: ChangeType) -> str:
        """Map change type to signal type."""
        mapping = {
            ChangeType.COMMENT: "decision",      # Reviewer feedback = decision input
            ChangeType.TRACKED_CHANGE: "progress",
            ChangeType.CONTENT_DRIFT: "status_change",
            ChangeType.STRUCTURAL: "decision",
        }
        return mapping.get(change_type, "raw")

    def _format_change_detail(self, change: ExternalChange) -> str:
        """Format change as human-readable detail."""
        if change.change_type == ChangeType.COMMENT:
            return f"Reviewer comment on {change.section}: {change.comment_text}"
        elif change.change_type == ChangeType.TRACKED_CHANGE:
            return f"Tracked change in {change.section}: '{change.original_text[:50]}...' → '{change.new_text[:50]}...'"
        elif change.change_type == ChangeType.CONTENT_DRIFT:
            return f"Content drift detected in {change.section}"
        elif change.change_type == ChangeType.STRUCTURAL:
            return f"Structural change: {change.section} modified"
        return "Document change detected"

    def extract_all(self):
        """Extract signals from all registered documents and synced folders.

        Includes:
        1. Registered PRD ↔ document mappings (MS 365, Google Docs, Box)
        2. Folder sync providers (Dropbox, Google Drive, etc.)

        Returns:
            Object with signals and errors attributes for CLI compatibility.
        """
        from dataclasses import dataclass, field

        @dataclass
        class ExtractionResult:
            signals: list = field(default_factory=list)
            errors: list = field(default_factory=list)

        result = ExtractionResult()

        # Extract from registered PRD ↔ document mappings
        try:
            result.signals = self.extract()
        except Exception as e:
            result.errors.append(f"Registry extraction: {e}")

        # Sync from folder providers (Dropbox, Google Drive, etc.)
        if self._sync_all_folders:
            try:
                folder_changes = self._sync_all_folders()
                for change in folder_changes:
                    # Convert folder sync changes to signals
                    if change.change_type.value in ("file_added", "file_modified"):
                        # The synced file is now in inbox/raw/docflow/{provider}/
                        local_path = Path(change.path)
                        if local_path.exists():
                            doc_signals = self.extract_local_doc(local_path)
                            result.signals.extend(doc_signals)
                    elif change.change_type.value == "file_deleted":
                        result.signals.append(Signal(
                            source="docflow_folder_sync",
                            timestamp=datetime.now(timezone.utc).isoformat(),
                            raw_text=f"File removed: {change.path}",
                            signal_type="progress",
                            detail=f"Remote file deleted: {change.metadata.get('remote_path', change.path)}",
                            confidence=0.8,
                            metadata={"change_type": "file_deleted", **change.metadata},
                        ))
            except Exception as e:
                result.errors.append(f"Folder sync: {e}")

        return result

    def extract_local_doc(
        self,
        doc_path: Path,
        correlator = None,
    ) -> list[Signal]:
        """Extract signals from a local Office document.

        Processes .docx, .xlsx, .pptx files from inbox/raw/docflow/.

        Args:
            doc_path: Path to the local document file
            correlator: Optional correlator for initiative matching

        Returns:
            List of Signal objects
        """
        signals = []

        if not doc_path.exists():
            return signals

        ext = doc_path.suffix.lower()

        # Use local DOCX provider for Word documents
        if ext == ".docx":
            provider = self._providers.get(DocFormat.LOCAL_DOCX)
            if provider:
                try:
                    content, changes = provider.fetch_content(str(doc_path))

                    # Create signals from extracted changes (comments, tracked changes)
                    for change in changes:
                        if change.change_type == ChangeType.COMMENT:
                            signal = Signal(
                                source="docflow_local",
                                timestamp=change.timestamp or datetime.now(timezone.utc).isoformat(),
                                raw_text=change.comment_text or "",
                                people=[change.author] if change.author else [],
                                initiatives=[],  # Could use correlator to match
                                signal_type="decision",
                                detail=f"Comment from {change.author or 'reviewer'}: {(change.comment_text or '')[:100]}",
                                confidence=0.8,
                                metadata={
                                    "file": doc_path.name,
                                    "change_type": "comment",
                                    "section": change.section,
                                },
                            )
                            signals.append(signal)
                        elif change.change_type == ChangeType.TRACKED_CHANGE:
                            signal = Signal(
                                source="docflow_local",
                                timestamp=change.timestamp or datetime.now(timezone.utc).isoformat(),
                                raw_text=f"{change.original_text or ''} → {change.new_text or ''}",
                                people=[change.author] if change.author else [],
                                initiatives=[],
                                signal_type="progress",
                                detail=f"Tracked change: '{(change.original_text or '')[:50]}' → '{(change.new_text or '')[:50]}'",
                                confidence=0.8,
                                metadata={
                                    "file": doc_path.name,
                                    "change_type": "tracked_change",
                                    "section": change.section,
                                },
                            )
                            signals.append(signal)

                    # If no changes found but content exists, create a raw signal
                    if not changes and content:
                        signals.append(Signal(
                            source="docflow_local",
                            timestamp=datetime.now(timezone.utc).isoformat(),
                            raw_text=content[:500] + ("..." if len(content) > 500 else ""),
                            signal_type="raw",
                            detail=f"Document content from {doc_path.name}",
                            confidence=0.5,
                            metadata={"file": doc_path.name, "type": "docx"},
                        ))

                except Exception as e:
                    signals.append(Signal(
                        source="docflow_local",
                        timestamp=datetime.now(timezone.utc).isoformat(),
                        raw_text=f"Error processing {doc_path.name}: {e}",
                        signal_type="blocker",
                        detail=f"Failed to extract from {doc_path.name}",
                        confidence=0.5,
                        metadata={"error": str(e)},
                    ))

        elif ext == ".xlsx":
            # Excel: extract comments from worksheets
            # For now, just note the file presence
            signals.append(Signal(
                source="docflow_local",
                timestamp=datetime.now(timezone.utc).isoformat(),
                raw_text=f"Excel file detected: {doc_path.name}",
                signal_type="raw",
                detail=f"Excel document {doc_path.name} ready for processing",
                confidence=0.5,
                metadata={"file": doc_path.name, "type": "excel"},
            ))

        elif ext == ".pptx":
            # PowerPoint: extract speaker notes and comments
            signals.append(Signal(
                source="docflow_local",
                timestamp=datetime.now(timezone.utc).isoformat(),
                raw_text=f"PowerPoint file detected: {doc_path.name}",
                signal_type="raw",
                detail=f"PowerPoint document {doc_path.name} ready for processing",
                confidence=0.5,
                metadata={"file": doc_path.name, "type": "powerpoint"},
            ))

        return signals


def register_prd(
    prd_id: str,
    md_path: str,
    external_uri: str,
    doc_format: str = "local_docx",
) -> None:
    """Register a PRD for DocFlow sync.

    Usage:
        neut docflow register --prd reactor-ops-log \\
            --md docs/requirements/prd_reactor-ops-log.md \\
            --uri "https://company.sharepoint.com/..." \\
            --format ms_word
    """
    extractor = DocFlowReviewExtractor()
    registry_path = extractor._registry_path

    # Load existing registry
    if registry_path.exists():
        registry = json.loads(registry_path.read_text())
    else:
        registry = {"prds": {}}

    # Add/update PRD
    registry["prds"][prd_id] = {
        "md_path": md_path,
        "external_uri": external_uri,
        "format": doc_format,
        "registered_at": datetime.now(timezone.utc).isoformat(),
    }

    # Save
    registry_path.parent.mkdir(parents=True, exist_ok=True)
    registry_path.write_text(json.dumps(registry, indent=2))
