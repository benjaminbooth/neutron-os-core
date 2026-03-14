"""PandocDocxProvider — generates .docx from markdown.

Extracts the generation logic from docs/_tools/md_to_docx.py.
Uses that script as an external tool, or falls back to pandoc directly.
"""

from __future__ import annotations

import subprocess
from pathlib import Path
from typing import Any

from ...factory import DocFlowFactory
from ..base import (
    GenerationProvider,
    GenerationOptions,
    GenerationResult,
)


class PandocDocxProvider(GenerationProvider):
    """Generate .docx files from markdown using md_to_docx.py or pandoc."""

    def __init__(self, config: dict[str, Any] | None = None):
        config = config or {}
        self.toc = config.get("toc", True)
        self.toc_depth = config.get("toc_depth", 3)
        self.reference_doc = config.get("reference_doc")
        self.mermaid_renderer = config.get("mermaid_renderer", "mermaid.ink")

        # Locate md_to_docx.py relative to repo
        self._md_to_docx = self._find_md_to_docx()

    @staticmethod
    def _find_md_to_docx() -> Path | None:
        """Find the md_to_docx.py script."""
        # Walk up to find repo root
        path = Path(__file__).resolve()
        while path != path.parent:
            candidate = path / "docs" / "_tools" / "md_to_docx.py"
            if candidate.exists():
                return candidate
            path = path.parent
        return None

    def generate(
        self, source_path: Path, output_path: Path, options: GenerationOptions
    ) -> GenerationResult:
        """Generate .docx from markdown source."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        warnings: list[str] = []

        if self._md_to_docx and self._md_to_docx.exists():
            # Use the custom md_to_docx.py converter
            result = self._generate_with_md_to_docx(source_path, output_path, options)
            if result:
                result = self._inject_footer(output_path, options, result)
                return result
            warnings.append("md_to_docx.py failed, falling back to pandoc")

        # Fall back to pandoc
        result = self._generate_with_pandoc(source_path, output_path, options, warnings)
        return self._inject_footer(output_path, options, result)

    def _generate_with_md_to_docx(
        self, source: Path, output: Path, options: GenerationOptions
    ) -> GenerationResult | None:
        """Use md_to_docx.py for generation."""
        cmd = [
            "python3",
            str(self._md_to_docx),
            str(source),
            str(output),
        ]

        if options.toc or self.toc:
            cmd.append("--toc")
        cmd.append("--bookmarks")

        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                cwd=str(source.parent),
                timeout=120,
            )

            if result.returncode == 0 and output.exists():
                return GenerationResult(
                    output_path=output,
                    format="docx",
                    size_bytes=output.stat().st_size,
                )
            return None

        except (subprocess.TimeoutExpired, FileNotFoundError):
            return None

    def _generate_with_pandoc(
        self,
        source: Path,
        output: Path,
        options: GenerationOptions,
        warnings: list[str],
    ) -> GenerationResult:
        """Use pandoc directly for generation."""
        cmd = ["pandoc", str(source), "-o", str(output)]

        if options.toc or self.toc:
            cmd.extend(["--toc", f"--toc-depth={options.toc_depth or self.toc_depth}"])

        if options.reference_doc or self.reference_doc:
            ref = options.reference_doc or self.reference_doc
            cmd.extend(["--reference-doc", ref])

        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                cwd=str(source.parent),
                timeout=60,
            )

            if result.returncode != 0:
                warnings.append(f"pandoc stderr: {result.stderr[:200]}")

            if not output.exists():
                raise RuntimeError(f"pandoc failed to generate {output}: {result.stderr}")

            return GenerationResult(
                output_path=output,
                format="docx",
                size_bytes=output.stat().st_size,
                warnings=warnings,
            )

        except FileNotFoundError:
            raise RuntimeError(
                "pandoc not found. Install with: brew install pandoc (macOS) "
                "or apt install pandoc (Linux)"
            )

    def _inject_footer(
        self, artifact_path: Path, options: GenerationOptions, result: GenerationResult
    ) -> GenerationResult:
        """Inject source URL, version, and publication date into document footer."""
        if not options.footer_metadata:
            return result

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            return result  # python-docx not available

        try:
            doc = Document(str(artifact_path))

            # Access the default section's footer
            section = doc.sections[0]
            footer = section.footer

            # Clear existing footer content if any
            for paragraph in footer.paragraphs:
                # Keep first paragraph, clear others
                if paragraph != footer.paragraphs[0]:
                    p_elem = paragraph._element
                    p_elem.getparent().remove(p_elem)

            # Use the first paragraph or create new one
            if footer.paragraphs:
                footer_para = footer.paragraphs[0]
            else:
                footer_para = footer.add_paragraph()

            # Build footer text
            footer_parts = []

            if "source_url" in options.footer_metadata:
                footer_parts.append(f"Source: {options.footer_metadata['source_url']}")

            if "version" in options.footer_metadata:
                footer_parts.append(f"Version: {options.footer_metadata['version']}")

            if "published_at" in options.footer_metadata:
                pub_date = options.footer_metadata["published_at"]
                # Format as YYYY-MM-DD HH:MM
                if "T" in pub_date:
                    pub_date = pub_date.split("T")[0]
                footer_parts.append(f"Published: {pub_date}")

            footer_text = " | ".join(footer_parts)

            # Clear and set footer paragraph
            footer_para.clear()
            if footer_text:
                run = footer_para.add_run(footer_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Save document with footer
            doc.save(str(artifact_path))

        except Exception as e:
            # Footer injection is best-effort; don't fail the whole publish
            result.warnings.append(f"Footer injection failed: {str(e)}")

        return result

    def rewrite_links(self, artifact_path: Path, link_map: dict[str, str]) -> None:
        """Rewrite internal links in a .docx file.

        Opens the docx, finds hyperlinks that reference .md files,
        and replaces them with published URLs from the link map.
        """
        if not link_map:
            return

        try:
            from docx import Document
        except ImportError:
            return  # python-docx not available

        doc = Document(str(artifact_path))
        modified = False

        # Walk all paragraphs and their hyperlink relationships
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for old_ref, new_url in link_map.items():
                    if old_ref in run.text:
                        run.text = run.text.replace(old_ref, new_url)
                        modified = True

        # Also check hyperlink relationships in the document part
        try:
            part = doc.part
            for rel in part.rels.values():
                if rel.is_external:
                    target = rel._target
                    for old_ref, new_url in link_map.items():
                        if old_ref in target:
                            rel._target = target.replace(old_ref, new_url)
                            modified = True
        except Exception:
            pass

        if modified:
            doc.save(str(artifact_path))

    def get_output_extension(self) -> str:
        return ".docx"

    def supports_watermark(self) -> bool:
        return True


# Self-register with factory
DocFlowFactory.register("generation", "pandoc-docx", PandocDocxProvider)
