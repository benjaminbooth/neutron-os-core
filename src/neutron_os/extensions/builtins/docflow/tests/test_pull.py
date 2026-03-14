"""Unit tests for the DocFlow pull functionality.

Tests the reverse flow: external storage → local .md updates.
"""

import pytest
from unittest.mock import patch

from neutron_os.extensions.builtins.docflow.config import DocFlowConfig, GitPolicy, ProviderConfig
from neutron_os.extensions.builtins.docflow.engine import DocFlowEngine
from neutron_os.extensions.builtins.docflow.state import DocumentState, PublicationRecord


@pytest.fixture
def pull_engine(tmp_path):
    """Create an engine configured for pull testing."""
    config = DocFlowConfig(
        git=GitPolicy(
            require_clean=False,
            require_pushed=False,
            publish_branches=["*"],
        ),
        generation=ProviderConfig(provider="pandoc-docx"),
        storage=ProviderConfig(
            provider="local",
            settings={"base_dir": str(tmp_path / "published")},
        ),
        notification=ProviderConfig(provider="terminal"),
        repo_root=tmp_path,
    )
    return DocFlowEngine(config)


@pytest.fixture
def published_doc(pull_engine, tmp_path):
    """Create and publish a document for pull testing."""
    source = tmp_path / "docs" / "test-pull.md"
    source.parent.mkdir(parents=True, exist_ok=True)
    source.write_text("# Test Document\n\nOriginal content.\n")

    record = pull_engine.publish(source, storage_override="local")
    return {
        "doc_id": "test-pull",
        "source_path": source,
        "record": record,
    }


class TestEnginePull:
    """Tests for DocFlowEngine.pull() method."""

    def test_pull_document_not_found(self, pull_engine):
        """Pull raises error for unknown doc_id."""
        with pytest.raises(ValueError, match="Document not found"):
            pull_engine.pull("nonexistent-doc")

    def test_pull_no_published_version(self, pull_engine, tmp_path):
        """Pull raises error when document has no published version."""
        # Create a document state with no published record
        doc_state = DocumentState(
            doc_id="unpublished",
            source_path="docs/unpublished.md",
            status="local",
        )
        pull_engine.state_store.update(doc_state)

        with pytest.raises(ValueError, match="no published version"):
            pull_engine.pull("unpublished")

    def test_pull_dry_run_no_changes(self, published_doc, pull_engine):
        """Dry run with no external changes returns no diff."""
        result = pull_engine.pull(published_doc["doc_id"], dry_run=True)

        # Since we just published, external and local should match (modulo format)
        # This tests the flow, actual diff depends on docx extraction
        assert "changed" in result
        assert "source_path" in result
        assert result["source_path"] == str(published_doc["source_path"])

    def test_pull_dry_run_does_not_modify(self, published_doc, pull_engine, tmp_path):
        """Dry run should not modify the local file."""
        source_path = published_doc["source_path"]
        original_content = source_path.read_text()
        original_mtime = source_path.stat().st_mtime

        pull_engine.pull(published_doc["doc_id"], dry_run=True)

        # File should be unchanged
        assert source_path.read_text() == original_content
        assert source_path.stat().st_mtime == original_mtime

    def test_pull_with_comments(self, published_doc, pull_engine):
        """Pull with include_comments extracts comments."""
        result = pull_engine.pull(
            published_doc["doc_id"],
            dry_run=True,
            include_comments=True,
        )

        assert "comments" in result
        assert isinstance(result["comments"], list)


class TestExtractText:
    """Tests for artifact text extraction."""

    def test_extract_text_from_markdown(self, pull_engine, tmp_path):
        """Extract text from a .md file."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Header\n\nParagraph content.\n")

        text = pull_engine._extract_text_from_artifact(md_file)

        assert "Header" in text
        assert "Paragraph" in text

    def test_extract_text_unsupported_format(self, pull_engine, tmp_path):
        """Unsupported format raises ValueError."""
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4 fake content")

        with pytest.raises(ValueError, match="Unsupported artifact format"):
            pull_engine._extract_text_from_artifact(pdf_file)

    def test_extract_text_from_docx(self, pull_engine, tmp_path):
        """Extract text from a .docx file."""
        pytest.importorskip("docx")
        from docx import Document

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("Test Heading", 0)
        doc.add_paragraph("Test paragraph content.")
        doc.save(str(docx_path))

        text = pull_engine._extract_text_from_artifact(docx_path)

        assert "Test Heading" in text
        assert "Test paragraph content" in text


class TestExtractComments:
    """Tests for comment extraction from artifacts."""

    def test_extract_comments_empty_docx(self, pull_engine, tmp_path):
        """Extract comments from docx without comments returns empty list."""
        pytest.importorskip("docx")
        from docx import Document

        docx_path = tmp_path / "no-comments.docx"
        doc = Document()
        doc.add_paragraph("No comments here.")
        doc.save(str(docx_path))

        comments = pull_engine._extract_comments_from_artifact(docx_path)

        assert comments == []

    def test_extract_comments_non_docx(self, pull_engine, tmp_path):
        """Extract comments from non-docx returns empty list."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# No comments in markdown\n")

        comments = pull_engine._extract_comments_from_artifact(md_file)

        assert comments == []


class TestPullWorkflow:
    """Integration tests for the complete pull workflow."""

    def test_pull_updates_local_file(self, pull_engine, tmp_path):
        """Pull with changes updates the local file."""
        # Set up: create state with a storage_id pointing to a modified doc
        docs_dir = tmp_path / "docs"
        docs_dir.mkdir(parents=True, exist_ok=True)
        source = docs_dir / "modified.md"
        source.write_text("# Original\n\nOld content.\n")

        # Create a "modified" docx in storage
        pytest.importorskip("docx")
        from docx import Document

        storage_dir = tmp_path / "published"
        storage_dir.mkdir(parents=True, exist_ok=True)
        modified_docx = storage_dir / "modified.docx"
        doc = Document()
        doc.add_heading("Modified", 0)
        doc.add_paragraph("New content from external edit.")
        doc.save(str(modified_docx))

        # Create state pointing to this storage
        doc_state = DocumentState(
            doc_id="modified",
            source_path="docs/modified.md",
            status="published",
            published=PublicationRecord(
                storage_id=str(modified_docx),  # Local storage uses path as ID
                url=f"file://{modified_docx}",
                version="v1",
                published_at="2026-02-24T00:00:00Z",
                commit_sha="abc123",
                generation_provider="pandoc-docx",
                storage_provider="local",
            ),
        )
        pull_engine.state_store.update(doc_state)

        # Mock download to copy the test docx to the requested path
        import shutil

        def mock_download(storage_id, local_path):
            shutil.copy(modified_docx, local_path)
            return local_path

        with patch.object(
            pull_engine._create_storage_provider("local").__class__,
            "download",
            side_effect=mock_download,
        ):
            result = pull_engine.pull("modified", dry_run=False)

        assert result["changed"] is True
        # Local file should be updated
        new_content = source.read_text()
        assert "Modified" in new_content or "New content" in new_content
