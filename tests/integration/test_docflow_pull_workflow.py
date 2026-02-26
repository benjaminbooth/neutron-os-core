"""Integration tests for the DocFlow pull workflow.

Tests the complete pull cycle: publish → external edit simulation → pull back.
These tests hit real filesystem operations but mock external storage.
"""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from tools.docflow.config import DocFlowConfig, GitPolicy, ProviderConfig
from tools.docflow.engine import DocFlowEngine
from tools.docflow.state import DocumentState, PublicationRecord


@pytest.fixture
def workflow_engine(tmp_path):
    """Create an engine for workflow testing with isolated state."""
    config = DocFlowConfig(
        git=GitPolicy(
            require_clean=False,
            require_pushed=False,
            publish_branches=["*"],
        ),
        generation=ProviderConfig(provider="pandoc-docx"),
        storage=ProviderConfig(
            provider="local",
            settings={"base_dir": str(tmp_path / "storage")},
        ),
        notification=ProviderConfig(provider="terminal"),
        repo_root=tmp_path,
    )
    engine = DocFlowEngine(config)
    return engine


class TestPublishPullRoundTrip:
    """Tests for the complete publish → pull round-trip."""

    def test_publish_then_pull_no_changes(self, workflow_engine, tmp_path):
        """Publish then immediately pull should show no changes."""
        # Create and publish a document
        source = tmp_path / "docs" / "roundtrip.md"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text("# Round Trip Test\n\nOriginal content.\n")

        record = workflow_engine.publish(source, storage_override="local")
        assert record is not None
        assert record.version == "v1"

        # Pull should show no changes (same content)
        result = workflow_engine.pull("roundtrip", dry_run=True)

        # Changed might be True due to format differences, but content similar
        assert "source_path" in result

    def test_publish_modify_pull_cycle(self, workflow_engine, tmp_path):
        """Complete cycle: publish → simulate external edit → pull."""
        pytest.importorskip("docx")
        from docx import Document

        # 1. Create and publish original document
        source = tmp_path / "docs" / "modify-test.md"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text("# Original Title\n\nOriginal paragraph.\n")

        record = workflow_engine.publish(source, storage_override="local")
        storage_path = Path(record.storage_id) if "file://" not in record.storage_id else Path(record.url.replace("file://", ""))

        # 2. Simulate external edit by modifying the .docx
        # Find the actual docx file in storage
        storage_dir = tmp_path / "storage"
        docx_files = list(storage_dir.glob("*.docx"))
        assert len(docx_files) == 1, f"Expected 1 docx, found {docx_files}"

        docx_path = docx_files[0]
        doc = Document(str(docx_path))
        # Add a new paragraph to simulate external edit
        doc.add_paragraph("Added by external reviewer.")
        doc.save(str(docx_path))

        # 3. Pull should detect changes
        result = workflow_engine.pull("modify-test", dry_run=True)

        assert result["changed"] is True
        assert result["diff"] is not None

    def test_pull_creates_missing_source_file(self, workflow_engine, tmp_path):
        """Pull creates source file if it doesn't exist locally."""
        pytest.importorskip("docx")
        from docx import Document

        # Create a docx file directly in storage
        storage_dir = tmp_path / "storage"
        storage_dir.mkdir(parents=True, exist_ok=True)
        docx_path = storage_dir / "new-doc.docx"
        doc = Document()
        doc.add_heading("New Document", 0)
        doc.add_paragraph("Created externally.")
        doc.save(str(docx_path))

        # Create state pointing to this storage
        doc_state = DocumentState(
            doc_id="new-doc",
            source_path="docs/new-doc.md",
            status="published",
            published=PublicationRecord(
                storage_id=str(docx_path),
                url=f"file://{docx_path}",
                version="v1",
                published_at="2026-02-24T00:00:00Z",
                commit_sha="abc123",
                generation_provider="pandoc-docx",
                storage_provider="local",
            ),
        )
        workflow_engine.state_store.update(doc_state)

        # Source file doesn't exist yet
        source_path = tmp_path / "docs" / "new-doc.md"
        assert not source_path.exists()

        # Mock storage download
        with patch.object(
            workflow_engine._create_storage_provider("local").__class__,
            "download",
            side_effect=lambda sid, path: path.write_bytes(docx_path.read_bytes()) or path,
        ):
            result = workflow_engine.pull("new-doc", dry_run=False)

        # Source file should now exist
        assert result["changed"] is True
        assert source_path.exists()


class TestMultiDocumentWorkflow:
    """Tests for workflows involving multiple documents."""

    def test_publish_multiple_then_status(self, workflow_engine, tmp_path):
        """Publish multiple docs and verify status reports all."""
        docs_dir = tmp_path / "docs"
        docs_dir.mkdir(parents=True, exist_ok=True)

        # Publish 3 documents
        for i in range(3):
            source = docs_dir / f"doc{i}.md"
            source.write_text(f"# Document {i}\n\nContent for doc {i}.\n")
            workflow_engine.publish(source, storage_override="local")

        # Check status
        statuses = workflow_engine.status()

        assert len(statuses) == 3
        doc_ids = {s.doc_id for s in statuses}
        assert doc_ids == {"doc0", "doc1", "doc2"}

    def test_draft_then_publish_workflow(self, workflow_engine, tmp_path):
        """Draft → review → publish workflow."""
        source = tmp_path / "docs" / "draft-workflow.md"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text("# Draft Document\n\nInitial draft content.\n")

        # 1. Publish as draft
        draft_record = workflow_engine.publish(source, storage_override="local", draft=True)
        assert draft_record.version == "v1"

        status = workflow_engine.status(source)
        assert len(status) == 1
        assert status[0].status == "draft"
        assert status[0].active_draft is not None

        # 2. Update content
        source.write_text("# Final Document\n\nRevised content after review.\n")

        # 3. Publish final
        final_record = workflow_engine.publish(source, storage_override="local", draft=False)
        assert final_record.version == "v2"

        status = workflow_engine.status(source)
        assert status[0].status == "published"
        assert status[0].published is not None


class TestVersioning:
    """Tests for document versioning through publish cycles."""

    def test_version_increments_on_republish(self, workflow_engine, tmp_path):
        """Version increments each time a document is republished."""
        source = tmp_path / "docs" / "versioned.md"
        source.parent.mkdir(parents=True, exist_ok=True)

        for i in range(5):
            source.write_text(f"# Version {i+1}\n\nContent revision {i+1}.\n")
            record = workflow_engine.publish(source, storage_override="local")
            assert record.version == f"v{i+1}"

    def test_version_tracks_across_draft_and_publish(self, workflow_engine, tmp_path):
        """Version increments across both draft and publish operations."""
        source = tmp_path / "docs" / "mixed-versions.md"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text("# Initial\n\nContent.\n")

        # v1 as draft
        r1 = workflow_engine.publish(source, storage_override="local", draft=True)
        assert r1.version == "v1"

        # v2 as draft
        source.write_text("# Updated\n\nNew content.\n")
        r2 = workflow_engine.publish(source, storage_override="local", draft=True)
        assert r2.version == "v2"

        # v3 as final publish
        source.write_text("# Final\n\nFinal content.\n")
        r3 = workflow_engine.publish(source, storage_override="local", draft=False)
        assert r3.version == "v3"


class TestLinkRegistry:
    """Tests for cross-document link registry integration."""

    def test_published_docs_register_links(self, workflow_engine, tmp_path):
        """Published documents are registered for link rewriting."""
        source = tmp_path / "docs" / "linked.md"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text("# Linked Document\n\nSee other docs.\n")

        workflow_engine.publish(source, storage_override="local")

        link_map = workflow_engine.registry.build_link_map()

        # Should have entries for the document
        assert any("linked" in k for k in link_map)

    def test_cross_doc_references_rewritten(self, workflow_engine, tmp_path):
        """Links between documents are rewritten in generated output."""
        docs_dir = tmp_path / "docs"
        docs_dir.mkdir(parents=True, exist_ok=True)

        # Publish first doc
        doc1 = docs_dir / "doc1.md"
        doc1.write_text("# Document 1\n\nStandalone content.\n")
        workflow_engine.publish(doc1, storage_override="local")

        # Second doc references first
        doc2 = docs_dir / "doc2.md"
        doc2.write_text("# Document 2\n\nSee [doc1](doc1.md) for details.\n")
        workflow_engine.publish(doc2, storage_override="local")

        # Both should be in registry
        link_map = workflow_engine.registry.build_link_map()
        assert len(link_map) >= 2


class TestErrorHandling:
    """Tests for error handling in workflows."""

    def test_pull_nonexistent_storage(self, workflow_engine, tmp_path):
        """Pull fails gracefully when storage is unavailable."""
        # Create state with fake storage reference
        doc_state = DocumentState(
            doc_id="missing-storage",
            source_path="docs/missing.md",
            status="published",
            published=PublicationRecord(
                storage_id="/nonexistent/path/file.docx",
                url="file:///nonexistent/path/file.docx",
                version="v1",
                published_at="2026-02-24T00:00:00Z",
                commit_sha="abc123",
                generation_provider="pandoc-docx",
                storage_provider="local",
            ),
        )
        workflow_engine.state_store.update(doc_state)

        with pytest.raises(Exception):
            workflow_engine.pull("missing-storage")

    def test_generate_invalid_markdown(self, workflow_engine, tmp_path):
        """Generate handles malformed markdown gracefully."""
        source = tmp_path / "bad.md"
        source.write_text("# Unclosed [link\n\n```\nunclosed code block")

        # Should still produce output (pandoc is lenient)
        output = workflow_engine.generate(source)
        assert output.exists()
